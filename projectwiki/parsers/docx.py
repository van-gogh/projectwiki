from __future__ import annotations

from pathlib import Path

from ..models import ParsedBlock


def parse_docx(path: Path) -> list[ParsedBlock]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse .docx files") from exc

    doc = docx.Document(path)
    blocks: list[ParsedBlock] = []
    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            blocks.append(ParsedBlock("docx_paragraph", text, {"paragraph": idx}, {"style": style}))
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            text = " | ".join(c for c in cells if c)
            if text:
                blocks.append(ParsedBlock("table_row", text, {"table": table_idx, "row": row_idx}, {}))
    return blocks
